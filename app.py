import streamlit as st
import psycopg2
import anthropic
import markdown
from html2docx import html2docx
from docx import Document
from datetime import datetime
import os
import json
import io
from dotenv import load_dotenv

load_dotenv()

st.title("🔍 Text-to-SQL Assistant")
st.caption("Ask questions about your database in plain English")

# Initialize session state
if "session_data" not in st.session_state:
    st.session_state.session_data = []
if "selected_db" not in st.session_state:
    st.session_state.selected_db = None
if "schema" not in st.session_state:
    st.session_state.schema = None
if "db_error" not in st.session_state:
    st.session_state.db_error = None

@st.cache_data
def get_databases():
    conn = psycopg2.connect(
        host=os.getenv("DB_HOST"),
        port=os.getenv("DB_PORT"),
        dbname="postgres",
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASSWORD")
    )
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("""
        SELECT datname FROM pg_database
        WHERE datistemplate = false AND datname != 'postgres'
        ORDER BY datname
    """)
    databases = [row[0] for row in cur.fetchall()]
    conn.close()
    return databases

@st.cache_data
def get_schema(db_name):
    conn = psycopg2.connect(
        host=os.getenv("DB_HOST"),
        port=os.getenv("DB_PORT"),
        dbname=db_name,
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASSWORD")
    )
    cur = conn.cursor()
    cur.execute("""
        SELECT table_name, column_name, data_type
        FROM information_schema.columns
        WHERE table_schema = 'public'
        ORDER BY table_name, ordinal_position
    """)
    schema = {}
    for table, column, dtype in cur.fetchall():
        if table not in schema:
            schema[table] = []
        schema[table].append(f"{column} ({dtype})")
    conn.close()
    return schema

# Startup check
try:
    databases = get_databases()
    st.session_state.db_error = None
except Exception as e:
    st.session_state.db_error = str(e)
    databases = []

if st.session_state.db_error:
    st.error(
        "**Unable to connect to PostgreSQL.**\n\n"
        "Please check that:\n"
        "- PostgreSQL is running\n"
        "- Your `.env` file has the correct credentials\n\n"
        f"Error details: `{st.session_state.db_error}`"
    )
    st.stop()

# Sidebar
with st.sidebar:
    st.header("Database")
    selected_db = st.selectbox(
        "Select a database",
        [None] + databases,
        format_func=lambda x: "Select a database..." if x is None else x
    )

    if selected_db != st.session_state.selected_db:
        st.session_state.selected_db = selected_db
        st.session_state.schema = get_schema(selected_db) if selected_db else None
        st.session_state.session_data = []

    st.divider()
    st.header("Available Tables")
    if st.session_state.schema:
        for table, columns in st.session_state.schema.items():
            with st.expander(table):
                for col in columns:
                    st.text(col)

# Main tabs
tab_ask, tab_history = st.tabs(["Ask a Question", "Session History"])

with tab_ask:
    with st.form("question_form"):
        question = st.text_input("Ask a question about your database:")
        submitted = st.form_submit_button("Ask", type="primary")

    if submitted and question and not st.session_state.schema:
        st.warning("Please select a database first.")

    if submitted and question and st.session_state.schema:
        schema = st.session_state.schema
        schema_text = ""
        for table, columns in schema.items():
            schema_text += f"\nTable: {table}\n"
            schema_text += "\n".join([f"  - {col}" for col in columns])
            schema_text += "\n"

        client = anthropic.Anthropic()

        # Step 1 — generate SQL
        with st.status("Generating SQL query from your question...", expanded=True) as status:
            st.write("Analyzing your question and database schema...")
            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                messages=[
                    {"role": "user", "content": f"""
                    You are a SQL expert. Given the following database schema,
                    write a PostgreSQL query to answer the user's question.
                    Return ONLY the SQL query, nothing else.

                    Database schema:
                    {schema_text}

                    User question: {question}
                    """}
                ]
            )
            status.update(label="SQL query generated!", state="complete")

        # Assign and clean SQL query
        sql_query = message.content[0].text.strip()
        if sql_query.startswith("```"):
            sql_query = sql_query.split("\n", 1)[1]
        if sql_query.endswith("```"):
            sql_query = sql_query.rsplit("```", 1)[0].strip()

        with st.expander("Generated SQL", expanded=True):
            st.code(sql_query, language="sql")

        # Step 2 — execute query
        try:
            conn = psycopg2.connect(
                host=os.getenv("DB_HOST"),
                port=os.getenv("DB_PORT"),
                dbname=st.session_state.selected_db,
                user=os.getenv("DB_USER"),
                password=os.getenv("DB_PASSWORD")
            )
            cur = conn.cursor()
            cur.execute(sql_query)
            results = cur.fetchall()
            column_names = [desc[0] for desc in cur.description]
            conn.close()

            if not results:
                st.warning("No results found for this question.")
            else:
                results_text = ", ".join(column_names) + "\n"
                for row in results:
                    results_text += str(row) + "\n"

                # Step 3 — interpret results
                with st.status("Interpreting query results...", expanded=True) as status:
                    st.write(f"Processing {len(results)} row(s) returned by the query...")
                    interpretation = client.messages.create(
                        model="claude-sonnet-4-6",
                        max_tokens=1024,
                        messages=[
                            {"role": "user", "content": f"""
                            The user asked: "{question}"

                            The SQL query returned these results:
                            {results_text}

                            Please provide a clear, concise answer in natural language.
                            No SQL, no technical details — just a clean human-readable answer.
                            Use only text, headings, and numbered or bullet lists.
                            Do NOT use markdown tables.
                            """}
                        ]
                    )
                    status.update(label="Answer ready!", state="complete")

                answer = interpretation.content[0].text
                st.markdown(answer)

                st.session_state.session_data.append({
                    "question": question,
                    "sql_query": sql_query,
                    "answer": answer,
                    "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M")
                })

        except Exception as e:
            st.error(f"Query failed: {e}")

with tab_history:
    if not st.session_state.session_data:
        st.info("No questions asked yet. Go to **Ask a Question** to get started.")
    else:
        st.header("Session History")
        for i, item in enumerate(st.session_state.session_data, 1):
            with st.expander(f"Q{i}: {item['question']}  —  {item['timestamp']}"):
                st.code(item["sql_query"], language="sql")
                st.markdown(item["answer"])

        st.divider()
        default_name = f"{st.session_state.selected_db}_{datetime.now().strftime('%Y-%m-%d')}"
        session_name = st.text_input("Session name", value=default_name)

        def build_docx_bytes():
            doc = Document()
            doc.add_heading("Text-to-SQL Session Report", level=1)
            doc.add_paragraph(f"Database: {st.session_state.selected_db}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            for i, item in enumerate(st.session_state.session_data, 1):
                doc.add_heading(f"Question {i}: {item['question']}", level=2)
                doc.add_heading("SQL Query", level=3)
                doc.add_paragraph(item["sql_query"])
                doc.add_heading("Answer", level=3)
                html_content = markdown.markdown(item["answer"])
                tmp_bytes = html2docx(html_content, title="")
                tmp_doc = Document(io.BytesIO(tmp_bytes.getvalue()))
                for element in tmp_doc.element.body:
                    doc.element.body.append(element)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        col_save, col_download = st.columns([1, 1])

        with col_save:
            if st.button("💾 Save Session to Disk"):
                os.makedirs(session_name, exist_ok=True)
                with open(f"{session_name}/{session_name}.json", "w", encoding="utf-8") as f:
                    json.dump({
                        "database": st.session_state.selected_db,
                        "date": datetime.now().strftime("%d/%m/%Y %H:%M"),
                        "questions": st.session_state.session_data
                    }, f, ensure_ascii=False, indent=2)
                buf = build_docx_bytes()
                with open(f"{session_name}/{session_name}.docx", "wb") as f:
                    f.write(buf.read())
                st.success(f"Session saved in `{session_name}/`")

        with col_download:
            buf = build_docx_bytes()
            st.download_button(
                label="📄 Download Report (.docx)",
                data=buf,
                file_name=f"{session_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
