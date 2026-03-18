import psycopg2
import anthropic
from dotenv import load_dotenv
from docx import Document
from datetime import datetime
import os
import json
import markdown
from html2docx import html2docx
import io

load_dotenv()

# Connect to the default postgres db to list available databases
_conn = psycopg2.connect(
    host=os.getenv("DB_HOST"),
    port=os.getenv("DB_PORT"),
    dbname="postgres",
    user=os.getenv("DB_USER"),
    password=os.getenv("DB_PASSWORD")
)
_conn.autocommit = True
_cur = _conn.cursor()
_cur.execute("""
    SELECT datname FROM pg_database
    WHERE datistemplate = false AND datname != 'postgres'
    ORDER BY datname
""")
databases = [row[0] for row in _cur.fetchall()]
_conn.close()

print("Available databases:")
for i, db in enumerate(databases, 1):
    print(f"  {i}. {db}")

while True:
    choice = input("\nEnter the database number or name: ").strip()
    if choice.isdigit() and 1 <= int(choice) <= len(databases):
        selected_db = databases[int(choice) - 1]
        break
    elif choice in databases:
        selected_db = choice
        break
    else:
        print("Invalid choice, please try again.")

# Connect to the selected database
conn = psycopg2.connect(
    host=os.getenv("DB_HOST"),
    port=os.getenv("DB_PORT"),
    dbname=selected_db,
    user=os.getenv("DB_USER"),
    password=os.getenv("DB_PASSWORD")
)
print(f"\nConnected to: {selected_db}\n")

# Read the database schema
cursor = conn.cursor()
cursor.execute("""
    SELECT table_name, column_name, data_type
    FROM information_schema.columns
    WHERE table_schema = 'public'
    ORDER BY table_name, ordinal_position
""")

schema = {}
for table, column, dtype in cursor.fetchall():
    if table not in schema:
        schema[table] = []
    schema[table].append(f"{column} ({dtype})")

schema_text = ""
for table, columns in schema.items():
    schema_text += f"\nTable: {table}\n"
    schema_text += "\n".join([f"  - {col}" for col in columns])
    schema_text += "\n"

print("Available tables:")
for table, columns in schema.items():
    print(f"\n  {table}")
    for col in columns:
        print(f"    - {col}")
print()

client = anthropic.Anthropic()
session_data = []

print("Text-to-SQL Assistant — type 'exit' to quit\n")

# Main loop — keep asking questions until user types 'exit'
while True:
    question = input("Ask a question about your database: ")

    if question.lower() == "exit":
        if not session_data:
            print("No questions asked, session not saved.")
            conn.close()
            exit()
        include_sql_answer = input("Include SQL queries in the report? (yes/no): ").strip().lower()
        include_sql_in_docx = include_sql_answer == "yes"
        break

    # Generate SQL
    print("\nGenerating SQL...")
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[
            {"role": "user", "content": f"""
            You are a SQL expert. Given the following database schema, 
            write a PostgreSQL query to answer the user's question.
            
            Return ONLY the SQL query, nothing else. No explanations, no markdown, just the raw SQL.
            
            Database schema:
            {schema_text}
            
            User question: {question}
            """}
        ]
    )

    sql_query = message.content[0].text
    print(f"\nGenerated SQL:\n{sql_query}")

    # Execute the query
    cursor = conn.cursor()
    try:
        cursor.execute(sql_query)
        results = cursor.fetchall()
        column_names = [desc[0] for desc in cursor.description]
    except Exception as e:
        conn.rollback()
        print(f"\nThe query failed to execute: {e}")
        print("Please try rephrasing your question.\n")
        continue

    if not results:
        print("\nNo results found for this question.\n")
        continue

    results_text = ", ".join(column_names) + "\n"
    for row in results:
        results_text += str(row) + "\n"

    # Interpret results
    print("\nInterpreting results...")
    interpretation = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[
            {"role": "user", "content": f"""
            The user asked: "{question}"
            
            The SQL query returned these results:
            {results_text}
            
            Please provide a clear, concise answer in natural language.
            No SQL, no technical details — just a clean human-readable answer.
            Use only text, headings, and numbered or bullet lists.
            Do NOT use markdown tables — represent data as numbered rankings or bullet points instead.
            Keep the answer concise and well structured.
            """}
        ]
    )

    answer = interpretation.content[0].text
    print(f"\nAnswer:\n{answer}\n")

    # Store in session
    session_data.append({
        "question": question,
        "sql_query": sql_query,
        "answer": answer,
        "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M")
    })

# Save session
default_name = f"{selected_db}_{datetime.now().strftime('%Y-%m-%d')}"
session_name = input(f"\nName this session (no extension) [{default_name}]: ").strip()
if not session_name:
    session_name = default_name

os.makedirs(session_name, exist_ok=True)

# Save JSON
with open(f"{session_name}/{session_name}.json", "w", encoding="utf-8") as f:
    json.dump({
        "database": selected_db,
        "date": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "questions": session_data
    }, f, ensure_ascii=False, indent=2)

# Save DOCX
doc = Document()
doc.add_heading("Text-to-SQL Session Report", level=1)
doc.add_paragraph(f"Database: {selected_db}")
doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

for i, item in enumerate(session_data, 1):
    doc.add_heading(f"Question {i}: {item['question']}", level=2)
    if include_sql_in_docx:
        doc.add_heading("SQL Query", level=3)
        doc.add_paragraph(item["sql_query"])
    doc.add_heading("Answer", level=3)
    html_content = markdown.markdown(item["answer"])
    tmp_bytes = html2docx(html_content, title="")
    tmp_doc = Document(io.BytesIO(tmp_bytes.getvalue()))
    for element in tmp_doc.element.body:
        doc.element.body.append(element)

doc.save(f"{session_name}/{session_name}.docx")

print(f"\nSession saved in {session_name}/")

conn.close()